from __future__ import annotations

from pathlib import Path

import docx

from epictrace.interfaces.media import MediaProcessor, MediaResult


class DocxMediaProcessor(MediaProcessor):
    def supports(self, path: Path) -> bool:
        return path.suffix.lower() == ".docx"

    def process(self, path: Path, *, progress_cb=None, cancel=None) -> MediaResult:
        document = docx.Document(str(path))
        text = "\n".join(p.text for p in document.paragraphs)
        return MediaResult(text=text, metadata={"paragraphs": len(document.paragraphs)})
